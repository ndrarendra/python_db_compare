"""
Web-enabled Database Comparison Tool with Exclude Tables Option,
Enhanced Excel & Word Reports, and Improved Loading Animation with Comments
--------------------------------------------------------------------------------

This tool compares two databases (MySQL/MariaDB) by checking their schema,
metadata, and data. It supports including or excluding specific tables via JSON.
Reports are generated as Excel and Word documents, and a web-based summary is shown.
"""

import os
import re
import zlib
import difflib
import datetime
import logging
from collections import defaultdict
from typing import Dict, List, Any, Optional, Tuple

import pymysql
import pandas as pd
from flask import Flask, request, render_template_string, send_from_directory

# Optional: for exporting Word documents
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

# ---------------------------
# Logging Setup Functions
# ---------------------------
class ReadableFormatter(logging.Formatter):
    def format(self, record):
        message = super().format(record)
        # Add extra newlines for special messages
        if record.msg.startswith(">>") or record.msg.startswith("=") or record.msg.startswith("----"):
            return "\n" + message
        return message

class ColoredFormatter(logging.Formatter):
    COLORS = {'DEBUG': '\033[94m','INFO': '\033[92m','WARNING': '\033[93m','ERROR': '\033[91m','CRITICAL': '\033[95m'}
    RESET = '\033[0m'
    def format(self, record):
        color = self.COLORS.get(record.levelname, self.RESET)
        message = super().format(record)
        return f"{color}{message}{self.RESET}"

def setup_logger(name: str, filename: str, verbose: bool = False) -> logging.Logger:
    """Sets up a logger to log messages to a file and optionally to the console."""
    logger = logging.getLogger(name)
    logger.setLevel(logging.INFO)
    file_handler = logging.FileHandler(filename, mode='w', encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    file_formatter = ReadableFormatter('[%(asctime)s] [%(levelname)s] [%(name)s:%(funcName)s:%(lineno)d] %(message)s')
    file_handler.setFormatter(file_formatter)
    logger.addHandler(file_handler)
    if verbose:
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        console_formatter = ColoredFormatter('[%(levelname)s] %(message)s')
        console_handler.setFormatter(console_formatter)
        logger.addHandler(console_handler)
    logger.propagate = False
    return logger



def compare_constraints(logger: logging.Logger,
                        constraintsA: Dict[Any, Dict[str, Any]],
                        constraintsB: Dict[Any, Dict[str, Any]],
                        dbA_name: str,
                        dbB_name: str,
                        summary: Dict[str, Any]) -> None:
    """
    Compares table constraints (e.g., primary keys, foreign keys, unique constraints)
    between two databases. Logs differences and updates the summary dictionary.

    :param logger: Logger object for logging comparison details.
    :param constraintsA: Constraints dictionary from the first database.
    :param constraintsB: Constraints dictionary from the second database.
    :param dbA_name: Name of the first database.
    :param dbB_name: Name of the second database.
    :param summary: Dictionary to store summary statistics of the comparison.
    """
    # Ensure the summary has a key for constraint differences
    summary.setdefault('constraints_differ', 0)
    
    # Get sets of constraint keys (each key is typically a tuple: (table_name, constraint_name))
    keysA = set(constraintsA.keys())
    keysB = set(constraintsB.keys())
    
    # Identify constraints unique to each database and those common to both
    only_in_A = keysA - keysB
    only_in_B = keysB - keysA
    common = keysA & keysB

    # Log header for constraint comparison
    logger.info("\n" + "="*50)
    logger.info(">> CONSTRAINT COMPARISON BETWEEN %s AND %s", dbA_name, dbB_name)
    logger.info("="*50)
    
    # Record counts of unique and common constraints in the summary
    summary['constraints_only_in_A'] = len(only_in_A)
    summary['constraints_only_in_B'] = len(only_in_B)
    summary['constraints_in_both'] = len(common)

    # Log any constraints that exist only in one of the databases
    if only_in_A:
        logger.info("Constraints only in %s: %s", dbA_name, sorted(only_in_A))
    if only_in_B:
        logger.info("Constraints only in %s: %s", dbB_name, sorted(only_in_B))

    # For constraints common to both, compare their definitions
    for key in common:
        if constraintsA[key] != constraintsB[key]:
            table_name, constraint_name = key
            logger.info("Difference in constraint '%s' on table '%s':", constraint_name, table_name)
            logger.info("  %s definition: %s", dbA_name, constraintsA[key])
            logger.info("  %s definition: %s", dbB_name, constraintsB[key])
            summary['constraints_differ'] += 1



# ---------------------------
# Output Folder Setup
# ---------------------------
def setup_output_folder(folder_name: Optional[str] = None) -> str:
    """Creates and returns the output folder (using a timestamp if none provided)."""
    if folder_name is None:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        folder_name = f"db_compare_output_{timestamp}"
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    return folder_name

# ---------------------------
# Database Connection Functions
# ---------------------------
def get_mysql_connection(host: str, user: str, password: str, db: str, port: int = 3306) -> pymysql.connections.Connection:
    """Creates and returns a connection to a MySQL/MariaDB database."""
    return pymysql.connect(host=host, user=user, password=password, db=db, port=port,
                           charset='utf8mb4', cursorclass=pymysql.cursors.Cursor)

def get_server_info(conn: pymysql.connections.Connection) -> Dict[str, str]:
    """Retrieves server version, default storage engine, and character set."""
    with conn.cursor() as cursor:
        cursor.execute("SELECT @@version")
        version = cursor.fetchone()[0]
        cursor.execute("SHOW VARIABLES LIKE 'default_storage_engine'")
        row = cursor.fetchone()
        default_engine = row[1] if row and len(row) > 1 else "Unknown"
        cursor.execute("SELECT @@character_set_server")
        default_encoding = cursor.fetchone()[0]
    program = "MariaDB" if "MariaDB" in version else "MySQL"
    return {"version": version, "default_engine": default_engine, "program": program,
            "default_encoding": default_encoding}

# ---------------------------
# Schema Retrieval Functions
# ---------------------------
def get_tables_and_columns(conn: pymysql.connections.Connection, db_name: str) -> Dict[str, List[Any]]:
    """Retrieves table names and their column definitions from the information_schema."""
    query = """
        SELECT TABLE_NAME, COLUMN_NAME, DATA_TYPE, IS_NULLABLE, COLUMN_DEFAULT, COLUMN_TYPE
        FROM information_schema.columns
        WHERE table_schema = %s
        ORDER BY TABLE_NAME, ORDINAL_POSITION;
    """
    with conn.cursor() as cursor:
        cursor.execute(query, (db_name,))
        rows = cursor.fetchall()
    schema = defaultdict(list)
    for table_name, col_name, data_type, is_nullable, col_default, col_type in rows:
        schema[table_name].append((col_name, data_type, is_nullable, col_default, col_type))
    return schema

def get_indexes(conn: pymysql.connections.Connection, db_name: str) -> Dict[Any, List[Any]]:
    """Retrieves indexes from the information_schema for a given database."""
    query = """
      SELECT TABLE_NAME, INDEX_NAME, SEQ_IN_INDEX, COLUMN_NAME, NON_UNIQUE, INDEX_TYPE 
      FROM information_schema.statistics
      WHERE table_schema = %s
      ORDER BY TABLE_NAME, INDEX_NAME, SEQ_IN_INDEX;
    """
    with conn.cursor() as cursor:
        cursor.execute(query, (db_name,))
        rows = cursor.fetchall()
    indexes = defaultdict(list)
    for table_name, index_name, seq, col_name, non_unique, index_type in rows:
        indexes[(table_name, index_name)].append((seq, col_name, non_unique, index_type))
    return indexes

def get_triggers(conn: pymysql.connections.Connection, db_name: str) -> Dict[str, Dict[str, Any]]:
    """Retrieves trigger definitions from the information_schema (or via SHOW TRIGGERS as fallback)."""
    triggers = {}
    try:
        query = """
            SELECT TRIGGER_NAME, EVENT_MANIPULATION, EVENT_OBJECT_TABLE, ACTION_TIMING, ACTION_STATEMENT 
            FROM information_schema.triggers
            WHERE trigger_schema = %s;
        """
        with conn.cursor() as cursor:
            cursor.execute(query, (db_name,))
            rows = cursor.fetchall()
        for trigger_name, event, table, timing, statement in rows:
            triggers[trigger_name] = {"event": event, "table": table, "timing": timing, "statement": statement}
    except Exception:
        try:
            with conn.cursor() as cursor:
                cursor.execute("SHOW TRIGGERS FROM `{}`".format(db_name))
                rows = cursor.fetchall()
            for row in rows:
                trigger_name = row[0]
                triggers[trigger_name] = {"event": row[1], "table": row[2], "timing": row[4], "statement": row[3]}
        except Exception:
            triggers = {}
    return triggers

def get_routines(conn: pymysql.connections.Connection, db_name: str) -> Dict[Any, Dict[str, Any]]:
    """Retrieves stored routines (procedures/functions) from the information_schema or mysql.proc as fallback."""
    routines = {}
    try:
        query = """
            SELECT ROUTINE_NAME, ROUTINE_TYPE, DATA_TYPE, DTD_IDENTIFIER, ROUTINE_DEFINITION, SQL_DATA_ACCESS
            FROM information_schema.routines
            WHERE routine_schema = %s;
        """
        with conn.cursor() as cursor:
            cursor.execute(query, (db_name,))
            rows = cursor.fetchall()
        for routine_name, routine_type, data_type, dtd_identifier, definition, sql_data_access in rows:
            routines[(routine_name, routine_type)] = {"data_type": data_type, "dtd_identifier": dtd_identifier,
                                                        "definition": definition, "sql_data_access": sql_data_access}
    except Exception:
        try:
            query = """
                SELECT name, type, returns, body, sql_data_access
                FROM mysql.proc
                WHERE db = %s;
            """
            with conn.cursor() as cursor:
                cursor.execute(query, (db_name,))
                rows = cursor.fetchall()
            for name, type_, returns, body, sql_data_access in rows:
                routines[(name, type_)] = {"data_type": returns, "dtd_identifier": None,
                                           "definition": body, "sql_data_access": sql_data_access}
        except Exception:
            routines = {}
    return routines

def get_constraints(conn: pymysql.connections.Connection, db_name: str) -> Dict[Any, Dict[str, Any]]:
    """Retrieves table constraints (e.g., primary keys, foreign keys, unique constraints) from the information_schema."""
    query = """
        SELECT tc.TABLE_NAME, tc.CONSTRAINT_NAME, tc.CONSTRAINT_TYPE,
               COALESCE(GROUP_CONCAT(DISTINCT kcu.COLUMN_NAME ORDER BY kcu.ORDINAL_POSITION SEPARATOR ', '), '') AS columns
        FROM information_schema.table_constraints tc
        LEFT JOIN information_schema.key_column_usage kcu
          ON tc.TABLE_SCHEMA = kcu.TABLE_SCHEMA 
             AND tc.TABLE_NAME = kcu.TABLE_NAME 
             AND tc.CONSTRAINT_NAME = kcu.CONSTRAINT_NAME
        WHERE tc.TABLE_SCHEMA = %s
        GROUP BY tc.TABLE_NAME, tc.CONSTRAINT_NAME, tc.CONSTRAINT_TYPE;
    """
    constraints = {}
    with conn.cursor() as cursor:
        cursor.execute(query, (db_name,))
        rows = cursor.fetchall()
        for table_name, constraint_name, constraint_type, columns in rows:
            columns_list = columns.split(', ') if columns else []
            constraints[(table_name, constraint_name)] = {"type": constraint_type, "columns": columns_list}
    return constraints

def get_views(conn: pymysql.connections.Connection, db_name: str) -> Dict[str, str]:
    """Retrieves view definitions from the information_schema or via SHOW CREATE VIEW as fallback."""
    views = {}
    try:
        query = """
            SELECT TABLE_NAME, VIEW_DEFINITION
            FROM information_schema.views
            WHERE table_schema = %s;
        """
        with conn.cursor() as cursor:
            cursor.execute(query, (db_name,))
            rows = cursor.fetchall()
        for view_name, definition in rows:
            views[view_name] = definition
    except Exception:
        try:
            with conn.cursor() as cursor:
                cursor.execute("SHOW FULL TABLES FROM `{}` WHERE TABLE_TYPE='VIEW'".format(db_name))
                rows = cursor.fetchall()
                view_names = [row[0] for row in rows]
            for view_name in view_names:
                with conn.cursor() as cursor:
                    cursor.execute("SHOW CREATE VIEW `{}`".format(view_name))
                    row = cursor.fetchone()
                    views[view_name] = row[1] if row and len(row) > 1 else ""
        except Exception:
            views = {}
    return views

def get_primary_key_columns(conn: pymysql.connections.Connection, db_name: str, table_name: str) -> List[str]:
    """Retrieves the primary key columns for a specified table."""
    query = """
        SELECT COLUMN_NAME
        FROM information_schema.key_column_usage
        WHERE table_schema = %s AND table_name = %s AND constraint_name = 'PRIMARY'
        ORDER BY ordinal_position;
    """
    with conn.cursor() as cursor:
        cursor.execute(query, (db_name, table_name))
        rows = cursor.fetchall()
    return [row[0] for row in rows]

def get_table_metadata(conn: pymysql.connections.Connection, db_name: str) -> Dict[str, Dict[str, Any]]:
    """Retrieves table metadata (engine, collation, comment, etc.) from the information_schema."""
    query = """
        SELECT TABLE_NAME, ENGINE, TABLE_COLLATION, TABLE_COMMENT, CREATE_TIME, UPDATE_TIME
        FROM information_schema.tables
        WHERE table_schema = %s;
    """
    with conn.cursor() as cursor:
        cursor.execute(query, (db_name,))
        rows = cursor.fetchall()
    metadata = {}
    for row in rows:
        table_name, engine, collation, comment, create_time, update_time = row
        metadata[table_name] = {"ENGINE": engine, "TABLE_COLLATION": collation, "TABLE_COMMENT": comment,
                                "CREATE_TIME": create_time, "UPDATE_TIME": update_time}
    return metadata

# ---------------------------
# Helper Functions (Comparison Helpers)
# ---------------------------
def check_db_name_in_definition(definition: str, expected_db_name: str) -> bool:
    """Checks if the SQL definition contains the expected database name."""
    if definition is None:
        return True
    expected = expected_db_name.replace("`", "").lower()
    pattern = re.compile(r'\b(?:from|join)\s+`([^`]+)`\s*\.\s*`([^`]+)`', flags=re.IGNORECASE)
    matches = pattern.findall(definition)
    if not matches:
        pattern_no_backticks = re.compile(r'\b(?:from|join)\s+([^\s`\.]+)\s*\.\s*([^\s`]+)', flags=re.IGNORECASE)
        matches = pattern_no_backticks.findall(definition)
    if not matches:
        return True
    for db_ref, _ in matches:
        if db_ref.lower() != expected:
            return False
    return True

def normalize_sql_definition(definition: str, db_name: str) -> str:
    """Normalizes an SQL definition by replacing the actual database name with a placeholder and lowercasing."""
    if definition is None:
        return ''
    pattern = re.compile(r'`?' + re.escape(db_name) + r'`?\.', flags=re.IGNORECASE)
    normalized = pattern.sub(r'<DATABASE>.', definition)
    normalized = " ".join(normalized.lower().split())
    return normalized

def is_long_text_column(index: int, columns_info: List[Any]) -> bool:
    """Determines if a column is 'long text' (e.g., BLOB or large VARCHAR)."""
    col = columns_info[index]
    data_type = col[1].lower()
    if data_type in {"blob", "longblob"}:
        return True
    if data_type == "varchar":
        m = re.search(r'varchar\((\d+)\)', col[4])
        if m:
            size = int(m.group(1))
            return size > 10000
    return False

def safe_decode(value: bytes) -> str:
    """Attempts to decode a byte string using several encodings."""
    for enc in ("utf-8", "cp949", "euc-kr"):
        try:
            return value.decode(enc)
        except UnicodeDecodeError:
            continue
    return value.decode("utf-8", errors="replace")

def diff_location_summary(textA: str, textB: str, max_blocks: int = 30, max_lines_per_block: int = 10) -> str:
    """Generates a summary of differences between two texts using difflib."""
    linesA = textA.splitlines()
    linesB = textB.splitlines()
    matcher = difflib.SequenceMatcher(None, linesA, linesB)
    blocks = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        block = []
        block.append(f"{tag.upper()} in A[{i1}:{i2}] vs B[{j1}:{j2}]:")
        block.append("A:")
        a_lines = linesA[i1:i2]
        if len(a_lines) > max_lines_per_block:
            block.extend(a_lines[:max_lines_per_block])
            block.append(f"... (and {len(a_lines) - max_lines_per_block} more lines)")
        else:
            block.extend(a_lines)
        block.append("B:")
        b_lines = linesB[j1:j2]
        if len(b_lines) > max_lines_per_block:
            block.extend(b_lines[:max_lines_per_block])
            block.append(f"... (and {len(b_lines) - max_lines_per_block} more lines)")
        else:
            block.extend(b_lines)
        blocks.append("\n".join(block))
    if not blocks:
        return "No differences."
    if len(blocks) > max_blocks:
        blocks = blocks[:max_blocks] + [f"... and {len(blocks) - max_blocks} more difference blocks"]
    return "\n\n".join(blocks)

def compute_table_checksum(conn: pymysql.connections.Connection, table_name: str, columns: List[str], logger: logging.Logger) -> Any:
    """Computes a checksum for a table by concatenating column values and applying CRC32."""
    if not columns:
        return None
    col_concat = ",'#',".join([f"COALESCE(CAST(`{col}` AS CHAR), '')" for col in columns])
    query = f"SELECT BIT_XOR(CRC32(CONCAT_WS('#',{col_concat}))) FROM `{table_name}`"
    try:
        with conn.cursor() as cursor:
            cursor.execute(query)
            result = cursor.fetchone()[0]
        return result
    except pymysql.err.OperationalError as e:
        logger.error("Could not compute checksum for table '%s': %s", table_name, e)
        return None
    except Exception as ex:
        logger.error("Unexpected error computing checksum for '%s': %s", table_name, ex)
        return None

# ---------------------------
# Comparison Functions for Schema/Objects
# ---------------------------
def compare_table_metadata(logger: logging.Logger,
                           metadataA: Dict[str, Dict[str, Any]],
                           metadataB: Dict[str, Dict[str, Any]],
                           dbA_name: str,
                           dbB_name: str,
                           summary: Dict[str, Any]) -> None:
    """Compares table metadata (ENGINE, collation, comment) between two databases."""
    summary.setdefault('metadata_differ', 0)
    summary.setdefault('metadata_diff_tables', [])
    common_tables = set(metadataA.keys()) & set(metadataB.keys())
    logger.info("\n" + "="*50)
    logger.info(">> TABLE METADATA COMPARISON BETWEEN %s AND %s", dbA_name, dbB_name)
    logger.info("="*50)
    for table in sorted(common_tables):
        metaA = metadataA[table]
        metaB = metadataB[table]
        diff_fields = []
        for field in ['ENGINE', 'TABLE_COLLATION', 'TABLE_COMMENT']:
            if metaA.get(field) != metaB.get(field):
                diff_fields.append(field)
        if diff_fields:
            logger.info("Metadata difference in table '%s':", table)
            for field in diff_fields:
                logger.info("  %s: %s vs %s", field, metaA.get(field), metaB.get(field))
            summary['metadata_differ'] += 1
            summary.setdefault('metadata_diff_tables', []).append(table)

def compare_schemas(logger: logging.Logger,
                    schemaA: Dict[str, Any],
                    schemaB: Dict[str, Any],
                    dbA_name: str,
                    dbB_name: str,
                    summary: Dict[str, Any],
                    table_summaries: Dict[str, Dict[str, Any]]) -> None:
    """Compares table schemas (column definitions) between two databases."""
    summary.setdefault('tables_only_in_A', 0)
    summary.setdefault('tables_only_in_B', 0)
    summary.setdefault('tables_in_both', 0)
    summary.setdefault('columns_compared', 0)
    summary.setdefault('columns_only_in_A', 0)
    summary.setdefault('columns_only_in_B', 0)
    summary.setdefault('schema_diff_tables', [])
    tablesA = set(schemaA.keys())
    tablesB = set(schemaB.keys())
    only_in_A = tablesA - tablesB
    only_in_B = tablesB - tablesA
    in_both = tablesA & tablesB
    logger.info("\n" + "="*50)
    logger.info(">> SCHEMA COMPARISON BETWEEN %s AND %s", dbA_name, dbB_name)
    logger.info("="*50)
    summary['tables_only_in_A'] = len(only_in_A)
    summary['tables_only_in_B'] = len(only_in_B)
    summary['tables_in_both'] = len(in_both)
    if only_in_A:
        logger.info("Tables only in %s (%d): %s", dbA_name, len(only_in_A), sorted(only_in_A))
    else:
        logger.info("No tables unique to %s", dbA_name)
    for table_name in only_in_A:
        table_summaries[table_name] = {
            'row_count_A': 0,
            'row_count_B': 0,
            'col_count_A': len(schemaA[table_name]),
            'col_count_B': 0,
            'rows_only_in_A': 0,
            'rows_only_in_B': 0,
            'rows_mismatched': 0,
            'checksum_A': None,
            'checksum_B': None
        }
    if only_in_B:
        logger.info("Tables only in %s (%d): %s", dbB_name, len(only_in_B), sorted(only_in_B))
    else:
        logger.info("No tables unique to %s", dbB_name)
    for table_name in only_in_B:
        table_summaries[table_name] = {
            'row_count_A': 0,
            'row_count_B': 0,
            'col_count_A': 0,
            'col_count_B': len(schemaB[table_name]),
            'rows_only_in_A': 0,
            'rows_only_in_B': 0,
            'rows_mismatched': 0,
            'checksum_A': None,
            'checksum_B': None
        }
    if in_both:
        logger.info("Tables in both (%d): %s", len(in_both), sorted(in_both))
    else:
        logger.info("No common tables between %s and %s", dbA_name, dbB_name)
    for tbl in sorted(in_both):
        colsA_set = set(schemaA[tbl])
        colsB_set = set(schemaB[tbl])
        summary['columns_compared'] += max(len(colsA_set), len(colsB_set))
        logger.info("Table '%s': %d columns in %s vs. %d columns in %s", tbl, len(schemaA[tbl]), dbA_name, len(schemaB[tbl]), dbB_name)
        diffA = colsA_set - colsB_set
        diffB = colsB_set - colsA_set
        if diffA or diffB:
            if diffA:
                logger.info("  -> Columns in '%s' only in %s: %s", tbl, dbA_name, sorted(diffA))
                summary['columns_only_in_A'] += len(diffA)
            if diffB:
                logger.info("  -> Columns in '%s' only in %s: %s", tbl, dbB_name, sorted(diffB))
                summary['columns_only_in_B'] += len(diffB)
            if tbl not in summary['schema_diff_tables']:
                summary['schema_diff_tables'].append(tbl)
        else:
            if schemaA[tbl] != schemaB[tbl]:
                logger.info("  -> Columns match but order differs for table '%s'", tbl)
                if tbl not in summary['schema_diff_tables']:
                    summary['schema_diff_tables'].append(tbl)
            else:
                logger.info("  -> Columns match for table '%s'", tbl)

def compare_indexes(logger: logging.Logger,
                    indexesA: Dict[Any, List[Any]],
                    indexesB: Dict[Any, List[Any]],
                    dbA_name: str,
                    dbB_name: str,
                    summary: Dict[str, Any]) -> None:
    """Compares indexes between two databases."""
    keysA = set(indexesA.keys())
    keysB = set(indexesB.keys())
    only_in_A = keysA - keysB
    only_in_B = keysB - keysA
    common = keysA & keysB
    logger.info("\n" + "="*50)
    logger.info(">> INDEX COMPARISON BETWEEN %s AND %s", dbA_name, dbB_name)
    logger.info("="*50)
    summary['indexes_only_in_A'] = len(only_in_A)
    summary['indexes_only_in_B'] = len(only_in_B)
    summary['indexes_in_both'] = len(common)
    if only_in_A:
        logger.info("Indexes only in %s: %s", dbA_name, sorted(only_in_A))
    if only_in_B:
        logger.info("Indexes only in %s: %s", dbB_name, sorted(only_in_B))
    for key in common:
        if indexesA[key] != indexesB[key]:
            table_name, index_name = key
            logger.info("Difference in index '%s' on table '%s':", index_name, table_name)
            logger.info("  %s: %s", dbA_name, indexesA[key])
            logger.info("  %s: %s", dbB_name, indexesB[key])
            summary['indexes_differ'] = summary.get('indexes_differ', 0) + 1

def compare_triggers(logger: logging.Logger,
                     triggersA: Dict[str, Dict[str, Any]],
                     triggersB: Dict[str, Dict[str, Any]],
                     dbA_name: str,
                     dbB_name: str,
                     summary: Dict[str, Any]) -> None:
    """Compares triggers between two databases."""
    keysA = set(triggersA.keys())
    keysB = set(triggersB.keys())
    only_in_A = keysA - keysB
    only_in_B = keysB - keysA
    common = keysA & keysB
    logger.info("\n" + "="*50)
    logger.info(">> TRIGGER COMPARISON BETWEEN %s AND %s", dbA_name, dbB_name)
    logger.info("="*50)
    summary['triggers_only_in_A'] = len(only_in_A)
    summary['triggers_only_in_B'] = len(only_in_B)
    summary['triggers_in_both'] = len(common)
    if only_in_A:
        logger.info("Triggers only in %s: %s", dbA_name, sorted(only_in_A))
    if only_in_B:
        logger.info("Triggers only in %s: %s", dbB_name, sorted(only_in_B))
    for trig in common:
        stmtA = triggersA[trig].get("statement", "")
        stmtB = triggersB[trig].get("statement", "")
        validA = check_db_name_in_definition(stmtA, dbA_name)
        validB = check_db_name_in_definition(stmtB, dbB_name)
        if not validA or not validB:
            logger.info("Trigger '%s' references an unexpected database name.", trig)
            if not validA:
                logger.info("  %s: %s", dbA_name, stmtA)
            if not validB:
                logger.info("  %s: %s", dbB_name, stmtB)
            summary["triggers_differ"] = summary.get("triggers_differ", 0) + 1
            continue
        normA = normalize_sql_definition(stmtA, dbA_name)
        normB = normalize_sql_definition(stmtB, dbB_name)
        if normA != normB:
            logger.info("Logic difference in trigger '%s':", trig)
            logger.info("  %s normalized: %s", dbA_name, normA)
            logger.info("  %s normalized: %s", dbB_name, normB)
            summary["triggers_differ"] = summary.get("triggers_differ", 0) + 1

def compare_routines(logger: logging.Logger,
                     routinesA: Dict[Any, Dict[str, Any]],
                     routinesB: Dict[Any, Dict[str, Any]],
                     dbA_name: str,
                     dbB_name: str,
                     summary: Dict[str, Any]) -> None:
    """Compares stored routines (procedures/functions) between two databases."""
    keysA = set(routinesA.keys())
    keysB = set(routinesB.keys())
    only_in_A = keysA - keysB
    only_in_B = keysB - keysA
    common = keysA & keysB
    logger.info("\n" + "="*50)
    logger.info(">> ROUTINE COMPARISON BETWEEN %s AND %s", dbA_name, dbB_name)
    logger.info("="*50)
    summary["routines_only_in_A"] = len(only_in_A)
    summary["routines_only_in_B"] = len(only_in_B)
    summary["routines_in_both"] = len(common)
    if only_in_A:
        logger.info("Routines only in %s: %s", dbA_name, sorted(only_in_A))
    if only_in_B:
        logger.info("Routines only in %s: %s", dbB_name, sorted(only_in_B))
    for key in common:
        routine_name, routine_type = key
        defA = routinesA[key].get("definition", "")
        defB = routinesB[key].get("definition", "")
        validA = check_db_name_in_definition(defA, dbA_name)
        validB = check_db_name_in_definition(defB, dbB_name)
        if not validA or not validB:
            logger.info("Routine '%s' [%s] references an unexpected database name.", routine_name, routine_type)
            if not validA:
                logger.info("  %s: %s", dbA_name, defA)
            if not validB:
                logger.info("  %s: %s", dbB_name, defB)
            summary["routines_differ"] = summary.get("routines_differ", 0) + 1
            continue
        normA = normalize_sql_definition(defA, dbA_name)
        normB = normalize_sql_definition(defB, dbB_name)
        if normA != normB:
            logger.info("Logic difference in routine '%s' [%s]:", routine_name, routine_type)
            logger.info("  %s normalized: %s", dbA_name, normA)
            logger.info("  %s normalized: %s", dbB_name, normB)
            summary["routines_differ"] = summary.get("routines_differ", 0) + 1

def compare_views(logger: logging.Logger,
                  viewsA: Dict[str, str],
                  viewsB: Dict[str, str],
                  dbA_name: str,
                  dbB_name: str,
                  summary: Dict[str, Any]) -> None:
    """Compares view definitions between two databases."""
    summary.setdefault('views_only_in_A', 0)
    summary.setdefault('views_only_in_B', 0)
    summary.setdefault('views_in_both', 0)
    summary.setdefault('views_differ', 0)
    keysA = set(viewsA.keys())
    keysB = set(viewsB.keys())
    only_in_A = keysA - keysB
    only_in_B = keysB - keysA
    common = keysA & keysB
    logger.info("\n" + "="*50)
    logger.info(">> VIEW COMPARISON BETWEEN %s AND %s", dbA_name, dbB_name)
    logger.info("="*50)
    summary['views_only_in_A'] = len(only_in_A)
    summary['views_only_in_B'] = len(only_in_B)
    summary['views_in_both'] = len(common)
    if only_in_A:
        logger.info("Views only in %s (%d): %s", dbA_name, len(only_in_A), sorted(only_in_A))
    else:
        logger.info("No views unique to %s.", dbA_name)
    if only_in_B:
        logger.info("Views only in %s (%d): %s", dbB_name, len(only_in_B), sorted(only_in_B))
    else:
        logger.info("No views unique to %s.", dbB_name)
    if common:
        logger.info("Views in both (%d): %s", len(common), sorted(common))
    else:
        logger.info("No common views between %s and %s.", dbA_name, dbB_name)
    for view in sorted(common):
        defA = viewsA[view]
        defB = viewsB[view]
        validA = check_db_name_in_definition(defA, dbA_name)
        validB = check_db_name_in_definition(defB, dbB_name)
        if not validA or not validB:
            normA = normalize_sql_definition(defA, dbA_name)
            normB = normalize_sql_definition(defB, dbB_name)
            if normA == normB:
                pattern = re.compile(r'\b(?:from|join)\s+`?([^`\s\.]+)`?\.', flags=re.IGNORECASE)
                matchA = pattern.search(defA)
                matchB = pattern.search(defB)
                actualA = matchA.group(1) if matchA else "Not specified"
                actualB = matchB.group(1) if matchB else "Not specified"
                logger.info("View '%s' logic is identical, but database names differ:", view)
                logger.info("  Actual database in first query: %s (expected: %s)", actualA, dbA_name)
                logger.info("  Actual database in second query: %s (expected: %s)", actualB, dbB_name)
            else:
                logger.info("View '%s' references unexpected database name(s), and logic differs.", view)
                if not validA:
                    logger.info("  Expected %s in view '%s', but got:\n%s", dbA_name, view, defA)
                if not validB:
                    logger.info("  Expected %s in view '%s', but got:\n%s", dbB_name, view, defB)
                summary['views_differ'] += 1
            continue
        normA = normalize_sql_definition(defA, dbA_name)
        normB = normalize_sql_definition(defB, dbB_name)
        if normA != normB:
            logger.info("Logic difference in view '%s':", view)
            logger.info("  %s normalized:\n%s", dbA_name, normA)
            logger.info("  %s normalized:\n%s", dbB_name, normB)
            summary['views_differ'] += 1

# ---------------------------
# Comparison Functions for Data
# ---------------------------
def get_row_key(row: tuple, pk_indices: List[int]) -> Any:
    """Extracts the primary key value(s) from a row."""
    if row is None:
        return None
    if len(pk_indices) == 1:
        return row[pk_indices[0]]
    else:
        return tuple(row[i] for i in pk_indices)

def compare_table_data(logger: logging.Logger,
                       connA: pymysql.connections.Connection,
                       connB: pymysql.connections.Connection,
                       dbA_name: str,
                       dbB_name: str,
                       table_name: str,
                       pk_cols: List[str],
                       summary: Dict[str, Any],
                       table_summaries: Dict[str, Dict[str, Any]],
                       columns_info: List[Any] = None) -> None:
    """Compares row-level data for a given table between two databases."""
    if not pk_cols:
        logger.info("Skipping data comparison for %s (no primary key).", table_name)
        summary['tables_skipped_no_pk'] = summary.get('tables_skipped_no_pk', 0) + 1
        return
    pk_order = ", ".join([f"`{col}`" for col in pk_cols])
    cursorA = connA.cursor()
    cursorB = connB.cursor()
    cursorA.execute(f"USE `{dbA_name}`")
    cursorB.execute(f"USE `{dbB_name}`")
    query = f"SELECT * FROM `{table_name}` ORDER BY {pk_order}"
    cursorA.execute(query)
    cursorB.execute(query)
    descA = cursorA.description
    col_names = [d[0] for d in descA]
    try:
        pk_indices = [col_names.index(pk) for pk in pk_cols]
    except ValueError as e:
        logger.error("Primary key column not found in table '%s': %s", table_name, e)
        return
    long_text_indices = []
    if columns_info:
        for i, _ in enumerate(columns_info):
            if is_long_text_column(i, columns_info):
                long_text_indices.append(i)
    def normalize_row(row: tuple) -> tuple:
        row_list = list(row)
        for i in long_text_indices:
            value = row_list[i]
            if value is not None:
                if isinstance(value, (bytes, bytearray)):
                    row_list[i] = zlib.crc32(value)
                elif isinstance(value, str):
                    row_list[i] = zlib.crc32(value.encode('utf-8'))
        return tuple(row_list)
    table_summaries.setdefault(table_name, {
        'row_count_A': 0,
        'row_count_B': 0,
        'col_count_A': len(columns_info) if columns_info else 0,
        'col_count_B': len(columns_info) if columns_info else 0,
        'rows_only_in_A': 0,
        'rows_only_in_B': 0,
        'rows_mismatched': 0,
        'checksum_A': None,
        'checksum_B': None
    })
    summary.setdefault('row_diff_details', {})
    countA = 0
    countB = 0
    rowA = cursorA.fetchone()
    rowB = cursorB.fetchone()
    while rowA is not None or rowB is not None:
        keyA = get_row_key(rowA, pk_indices)
        keyB = get_row_key(rowB, pk_indices)
        if rowA is not None and rowB is not None:
            if keyA == keyB:
                countA += 1
                countB += 1
                normA = normalize_row(rowA)
                normB = normalize_row(rowB)
                if normA != normB:
                    logger.info("---- Data mismatch in table '%s' for PK=%s ----", table_name, keyA)
                    for i, col_name in enumerate(col_names):
                        valA = rowA[i]
                        valB = rowB[i]
                        if valA != valB:
                            if columns_info and is_long_text_column(i, columns_info):
                                strA = safe_decode(valA) if isinstance(valA, bytes) else str(valA)
                                strB = safe_decode(valB) if isinstance(valB, bytes) else str(valB)
                                diff_summary = diff_location_summary(strA, strB)
                                logger.info("Column '%s' differences (detailed):\n%s", col_name, diff_summary)
                            else:
                                logger.info("Column '%s' differs. %s: %s | %s: %s", col_name, dbA_name, valA, dbB_name, valB)
                    summary['rows_mismatched'] = summary.get('rows_mismatched', 0) + 1
                    table_summaries[table_name]['rows_mismatched'] += 1
                    summary.setdefault('row_diff_details', {}).setdefault(table_name, []).append(str(keyA))
                rowA = cursorA.fetchone()
                rowB = cursorB.fetchone()
            elif keyA < keyB:
                logger.info("---- Row with PK=%s exists only in %s ----", keyA, dbA_name)
                summary['rows_only_in_A'] = summary.get('rows_only_in_A', 0) + 1
                table_summaries[table_name]['rows_only_in_A'] += 1
                countA += 1
                summary.setdefault('row_diff_details', {}).setdefault(table_name, []).append(str(keyA))
                rowA = cursorA.fetchone()
            else:
                logger.info("---- Row with PK=%s exists only in %s ----", keyB, dbB_name)
                summary['rows_only_in_B'] = summary.get('rows_only_in_B', 0) + 1
                table_summaries[table_name]['rows_only_in_B'] += 1
                countB += 1
                summary.setdefault('row_diff_details', {}).setdefault(table_name, []).append(str(keyB))
                rowB = cursorB.fetchone()
        elif rowA is not None:
            keyA = get_row_key(rowA, pk_indices)
            logger.info("---- Row with PK=%s exists only in %s ----", keyA, dbA_name)
            summary['rows_only_in_A'] = summary.get('rows_only_in_A', 0) + 1
            table_summaries[table_name]['rows_only_in_A'] += 1
            countA += 1
            summary.setdefault('row_diff_details', {}).setdefault(table_name, []).append(str(keyA))
            rowA = cursorA.fetchone()
        elif rowB is not None:
            keyB = get_row_key(rowB, pk_indices)
            logger.info("---- Row with PK=%s exists only in %s ----", keyB, dbB_name)
            summary['rows_only_in_B'] = summary.get('rows_only_in_B', 0) + 1
            table_summaries[table_name]['rows_only_in_B'] += 1
            countB += 1
            summary.setdefault('row_diff_details', {}).setdefault(table_name, []).append(str(keyB))
            rowB = cursorB.fetchone()
    table_summaries[table_name]['row_count_A'] = countA
    table_summaries[table_name]['row_count_B'] = countB
    summary['total_rows_compared'] = summary.get('total_rows_compared', 0) + countA + countB
    cursorA.close()
    cursorB.close()

def compare_all_tables_data(logger: logging.Logger,
                            connA: pymysql.connections.Connection,
                            connB: pymysql.connections.Connection,
                            dbA_name: str,
                            dbB_name: str,
                            schemaA: Dict[str, List[Any]],
                            schemaB: Dict[str, List[Any]],
                            summary: Dict[str, Any],
                            table_summaries: Dict[str, Dict[str, Any]]) -> None:
    """Compares row-level data for all common tables."""
    summary.setdefault('total_tables_compared', 0)
    summary.setdefault('total_rows_compared', 0)
    tablesA = set(schemaA.keys())
    tablesB = set(schemaB.keys())
    common_tables = tablesA & tablesB
    logger.info("\n" + "="*50)
    logger.info(">> DATA COMPARISON BETWEEN %s AND %s", dbA_name, dbB_name)
    logger.info("="*50)
    summary['total_tables_compared'] = len(common_tables)
    for table_name in sorted(common_tables):
        try:
            with connA.cursor() as curA, connB.cursor() as curB:
                curA.execute(f"SELECT COUNT(*) FROM `{table_name}`")
                countA = curA.fetchone()[0] if curA.rowcount > 0 else 0
                curB.execute(f"SELECT COUNT(*) FROM `{table_name}`")
                countB = curB.fetchone()[0] if curB.rowcount > 0 else 0
            logger.info("Table '%s': %s=%d rows, %s=%d rows", table_name, dbA_name, countA, dbB_name, countB)
            col_info_A = schemaA[table_name]
            columnsA = {col[0] for col in col_info_A}
            columnsB = {col[0] for col in schemaB[table_name]}
            common_cols = list(columnsA & columnsB)
            checksumA = compute_table_checksum(connA, table_name, common_cols, logger)
            checksumB = compute_table_checksum(connB, table_name, common_cols, logger)
            table_summaries[table_name] = {
                'row_count_A': countA,
                'row_count_B': countB,
                'col_count_A': len(col_info_A),
                'col_count_B': len(schemaB[table_name]),
                'rows_only_in_A': 0,
                'rows_only_in_B': 0,
                'rows_mismatched': 0,
                'checksum_A': checksumA,
                'checksum_B': checksumB
            }
            summary['total_rows_compared'] += countA + countB
            pk_cols = get_primary_key_columns(connA, dbA_name, table_name)
            compare_table_data(logger, connA, connB, dbA_name, dbB_name, table_name, pk_cols, summary, table_summaries, col_info_A)
        except Exception as e:
            logger.error("Error comparing data for table '%s': %s", table_name, e)
            table_summaries[table_name] = {
                'row_count_A': 0,
                'row_count_B': 0,
                'col_count_A': 0,
                'col_count_B': 0,
                'rows_only_in_A': 0,
                'rows_only_in_B': 0,
                'rows_mismatched': 0,
                'checksum_A': None,
                'checksum_B': None,
                'error': str(e)
            }
            continue

def generate_brief_summary(summary: Dict[str, Any]) -> str:
    """Generates a multi-line summary report of the comparison results."""
    dbA_info = summary.get('dbA_info', {})
    dbB_info = summary.get('dbB_info', {})
    brief = []
    brief.append("----- Database Information -----")
    brief.append(f"Database A: Name: {summary.get('dbA_name', 'Unknown')}, Program: {dbA_info.get('program', 'Unknown')}, Version: {dbA_info.get('version', 'Unknown')}, Default Engine: {dbA_info.get('default_engine', 'Unknown')}, Main Encoding: {dbA_info.get('default_encoding', 'Unknown')}")
    brief.append(f"Database B: Name: {summary.get('dbB_name', 'Unknown')}, Program: {dbB_info.get('program', 'Unknown')}, Version: {dbB_info.get('version', 'Unknown')}, Default Engine: {dbB_info.get('default_engine', 'Unknown')}, Main Encoding: {dbB_info.get('default_encoding', 'Unknown')}")
    brief.append("--------------------------------")
    total_common_tables = summary.get('tables_in_both', 0)
    schema_diff_count = len(summary.get('schema_diff_tables', []))
    schema_match_percent = 100.0 if total_common_tables == 0 else ((total_common_tables - schema_diff_count) / total_common_tables * 100)
    total_rows = summary.get('total_rows_compared', 0)
    unique_rows = total_rows / 2 if total_rows else 0
    rows_mismatched = summary.get('rows_mismatched', 0)
    row_match_percent = 100.0 if unique_rows == 0 else ((unique_rows - rows_mismatched) / unique_rows * 100)
    metadata_diff_count = len(summary.get('metadata_diff_tables', []))
    metadata_match_percent = 100.0 if total_common_tables == 0 else ((total_common_tables - metadata_diff_count) / total_common_tables * 100)
    brief.append(f"Schema Match: {schema_match_percent:.2f}% of common tables have matching column definitions.")
    brief.append(f"Metadata Match: {metadata_match_percent:.2f}% of common tables have matching metadata.")
    brief.append(f"Data Match: {row_match_percent:.2f}% of rows match across common tables.")
    if summary.get('schema_diff_tables'):
        brief.append("Tables with column differences:")
        for tbl in summary['schema_diff_tables']:
            brief.append(f"  - {tbl}")
    else:
        brief.append("No differences in table columns detected.")
    if summary.get('metadata_diff_tables'):
        brief.append("Tables with metadata differences:")
        for tbl in summary['metadata_diff_tables']:
            brief.append(f"  - {tbl}")
    else:
        brief.append("No differences in table metadata detected.")
    if summary.get('row_diff_details'):
        brief.append("Tables with row differences (primary key values):")
        for tbl, keys in summary['row_diff_details'].items():
            brief.append(f"  - {tbl}: {', '.join(keys)}")
    else:
        brief.append("No row differences detected.")
    brief.append("--------------------------------")
    return "\n".join(brief)

# ---------------------------
# Excel Report Generator (Enhanced)
# ---------------------------
def export_table_summary_to_excel(table_summaries: Dict[str, Dict[str, Any]], filename: str, dbA_name: str, dbB_name: str) -> None:
    """Generates an Excel summary report with conditional formatting for the Row Diff(%) column."""
    data = []
    counter = 1
    for table, summ_data in sorted(table_summaries.items()):
        if summ_data.get('col_count_A', 0) > 0 and summ_data.get('col_count_B', 0) == 0:
            diff_text = f"Only in {dbA_name}"
        elif summ_data.get('col_count_B', 0) > 0 and summ_data.get('col_count_A', 0) == 0:
            diff_text = f"Only in {dbB_name}"
        else:
            row_count_A = summ_data.get('row_count_A', 0)
            row_count_B = summ_data.get('row_count_B', 0)
            max_rows = max(row_count_A, row_count_B) if max(row_count_A, row_count_B) > 0 else 1
            diff_value = (summ_data.get('rows_only_in_A', 0) + summ_data.get('rows_only_in_B', 0) +
                          summ_data.get('rows_mismatched', 0)) / max_rows * 100
            diff_text = "Same" if abs(diff_value) < 1e-9 else f"{diff_value:.2f}% diff"
        row = {
            'No.': counter,
            'Table': table,
            f'col_count_{dbA_name}': summ_data.get('col_count_A', 0),
            f'col_count_{dbB_name}': summ_data.get('col_count_B', 0),
            f'row_count_{dbA_name}': summ_data.get('row_count_A', 0),
            f'row_count_{dbB_name}': summ_data.get('row_count_B', 0),
            f'rows_only_in_{dbA_name}': summ_data.get('rows_only_in_A', 0),
            f'rows_only_in_{dbB_name}': summ_data.get('rows_only_in_B', 0),
            'rows_mismatched': summ_data.get('rows_mismatched', 0),
            'col_mismatch': abs(summ_data.get('col_count_A', 0) - summ_data.get('col_count_B', 0)),
            f'checksum_{dbA_name}': summ_data.get('checksum_A', ""),
            f'checksum_{dbB_name}': summ_data.get('checksum_B', ""),
            'Row Diff(%)': diff_text
        }
        data.append(row)
        counter += 1
    df = pd.DataFrame(data)
    with pd.ExcelWriter(filename, engine='xlsxwriter') as writer:
        startrow = 1
        df.to_excel(writer, index=False, sheet_name='Summary', startrow=startrow)
        workbook = writer.book
        worksheet = writer.sheets['Summary']
        title_format = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center', 'valign': 'vcenter'})
        worksheet.merge_range(0, 0, 0, len(df.columns)-1, "Database Comparison Summary", title_format)
        worksheet.freeze_panes(2, 0)
        for i, col in enumerate(df.columns):
            col_width = max(df[col].astype(str).map(len).max(), len(col)) + 4
            worksheet.set_column(i, i, col_width)
        diff_col_index = len(df.columns) - 1
        green_format = workbook.add_format({'bg_color': '#c6efce', 'font_color': '#006100'})
        red_format = workbook.add_format({'bg_color': '#ffc7ce', 'font_color': '#9c0006'})
        worksheet.conditional_format(startrow+1, diff_col_index, startrow+len(df), diff_col_index,
                                     {'type': 'cell', 'criteria': '==', 'value': '"Same"', 'format': green_format})
        worksheet.conditional_format(startrow+1, diff_col_index, startrow+len(df), diff_col_index,
                                     {'type': 'cell', 'criteria': '!=', 'value': '"Same"', 'format': red_format})

# ---------------------------
# Word Report Generator (Enhanced)
# ---------------------------
def export_summary_to_word(summary: Dict[str, Any],
                           table_summaries: Dict[str, Dict[str, Any]],
                           filename: str,
                           dbA_name: str,
                           dbB_name: str) -> None:
    """Generates a Word summary report with improved table formatting."""
    if Document is None:
        return
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(2)
    cover_title = document.add_heading("Database Comparison Report", 0)
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Company Name / Confidential", style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(datetime.datetime.now().strftime("%B %d, %Y"), style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    document.add_heading("1. Executive Summary", level=1)
    p = document.add_paragraph()
    p.add_run("This report presents a high-level comparison of two databases. ").bold = True
    p.add_run("Below is a summary of the schema, metadata, and data comparisons:")
    document.add_paragraph(generate_brief_summary(summary), style='Intense Quote')
    document.add_heading("2. Detailed Results", level=1)
    document.add_heading("2.1 Databases in Scope", level=2)
    table_db = document.add_table(rows=2, cols=3)
    table_db.style = "Light Shading Accent 2"
    hdr_cells = table_db.rows[0].cells
    hdr_cells[0].text = "Database Name"
    hdr_cells[1].text = "Program / Version"
    hdr_cells[2].text = "Encoding / Engine"
    row_cells = table_db.rows[1].cells
    row_cells[0].text = f"{dbA_name}\n{dbB_name}"
    row_cells[1].text = f"{summary.get('dbA_info', {}).get('program', 'Unknown')} {summary.get('dbA_info', {}).get('version', '')}\n{summary.get('dbB_info', {}).get('program', 'Unknown')} {summary.get('dbB_info', {}).get('version', '')}"
    row_cells[2].text = f"{summary.get('dbA_info', {}).get('default_encoding', 'Unknown')} / {summary.get('dbA_info', {}).get('default_engine', 'Unknown')}\n{summary.get('dbB_info', {}).get('default_encoding', 'Unknown')} / {summary.get('dbB_info', {}).get('default_engine', 'Unknown')}"
    document.add_paragraph()
    document.add_heading("2.2 Table-wise Summary", level=2)
    num_cols = 12
    word_table = document.add_table(rows=1, cols=num_cols)
    word_table.style = "Colorful List Accent 2"
    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = word_table.rows[0].cells
    headers = [
        "Table",
        f"col_count_{dbA_name}",
        f"col_count_{dbB_name}",
        f"row_count_{dbA_name}",
        f"row_count_{dbB_name}",
        f"rows_only_in_{dbA_name}",
        f"rows_only_in_{dbB_name}",
        "rows_mismatched",
        "col_mismatch",
        f"checksum_{dbA_name}",
        f"checksum_{dbB_name}",
        "Row Diff(%)"
    ]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for tname, data in sorted(table_summaries.items()):
        row_cells = word_table.add_row().cells
        row_cells[0].text = str(tname)
        row_cells[1].text = str(data.get('col_count_A', 0))
        row_cells[2].text = str(data.get('col_count_B', 0))
        row_cells[3].text = str(data.get('row_count_A', 0))
        row_cells[4].text = str(data.get('row_count_B', 0))
        row_cells[5].text = str(data.get('rows_only_in_A', 0))
        row_cells[6].text = str(data.get('rows_only_in_B', 0))
        row_cells[7].text = str(data.get('rows_mismatched', 0))
        row_cells[8].text = str(abs(data.get('col_count_A', 0) - data.get('col_count_B', 0)))
        row_cells[9].text = str(data.get('checksum_A', ""))
        row_cells[10].text = str(data.get('checksum_B', ""))
        if data.get('col_count_A', 0) > 0 and data.get('col_count_B', 0) == 0:
            diff_text = f"Only in {dbA_name}"
        elif data.get('col_count_B', 0) > 0 and data.get('col_count_A', 0) == 0:
            diff_text = f"Only in {dbB_name}"
        else:
            row_count_A = data.get('row_count_A', 0)
            row_count_B = data.get('row_count_B', 0)
            max_rows = max(row_count_A, row_count_B) if max(row_count_A, row_count_B) > 0 else 1
            diff_value = (data.get('rows_only_in_A', 0) + data.get('rows_only_in_B', 0) + data.get('rows_mismatched', 0)) / max_rows * 100
            diff_text = "Same" if abs(diff_value) < 1e-9 else f"{diff_value:.2f}% diff"
        row_cells[11].text = diff_text
    for row in word_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    try:
        for cell in word_table.columns[-1].cells:
            cell.width = Inches(1.5)
    except Exception:
        pass
    document.add_heading("3. Conclusion", level=1)
    document.add_paragraph("Based on the analysis above, the following issues were identified:")
    document.add_paragraph(generate_brief_summary(summary), style='Intense Quote')
    document.add_paragraph("For further information or to resolve any discrepancies, please refer to the detailed logs or contact the DBA team.")
    document.save(filename)

def generate_summary_log(summary_logger: logging.Logger,
                         summary: Dict[str, Any],
                         table_summaries: Dict[str, Dict[str, Any]]) -> None:
    """Logs a detailed summary of the comparison results."""
    brief = generate_brief_summary(summary)
    summary_logger.info("\n" + "="*50)
    summary_logger.info(">> DATABASE COMPARISON SUMMARY")
    summary_logger.info("="*50)
    summary_logger.info(brief)
    summary_logger.info("\nSchema Comparison:")
    summary_logger.info("  - Tables only in DB A: %d", summary.get('tables_only_in_A', 0))
    summary_logger.info("  - Tables only in DB B: %d", summary.get('tables_only_in_B', 0))
    summary_logger.info("  - Tables in both DBs: %d", summary.get('tables_in_both', 0))
    summary_logger.info("  - Columns compared: %d", summary.get('columns_compared', 0))
    summary_logger.info("  - Columns only in DB A: %d", summary.get('columns_only_in_A', 0))
    summary_logger.info("  - Columns only in DB B: %d", summary.get('columns_only_in_B', 0))
    summary_logger.info("  - Indexes differing: %d", summary.get('indexes_differ', 0))
    summary_logger.info("  - Triggers differing: %d", summary.get('triggers_differ', 0))
    summary_logger.info("  - Routines differing: %d", summary.get('routines_differ', 0))
    summary_logger.info("  - Constraints differing: %d", summary.get('constraints_differ', 0))
    summary_logger.info("  - Views differing: %d", summary.get('views_differ', 0))
    summary_logger.info("\nMetadata Comparison:")
    summary_logger.info("  - Tables with metadata differences: %d", len(summary.get('metadata_diff_tables', [])))
    if summary.get('metadata_diff_tables'):
        summary_logger.info("    Affected tables: %s", ", ".join(summary.get('metadata_diff_tables', [])))
    summary_logger.info("\nData Comparison:")
    summary_logger.info("  - Total tables compared: %d", summary.get('total_tables_compared', 0))
    summary_logger.info("  - Total rows compared: %d", summary.get('total_rows_compared', 0))
    summary_logger.info("  - Tables skipped (no primary key): %d", summary.get('tables_skipped_no_pk', 0))
    summary_logger.info("  - Rows only in DB A: %d", summary.get('rows_only_in_A', 0))
    summary_logger.info("  - Rows only in DB B: %d", summary.get('rows_only_in_B', 0))
    summary_logger.info("  - Rows mismatched: %d", summary.get('rows_mismatched', 0))
    summary_logger.info("="*50)
    summary_logger.info("\n--- Table-wise Summary ---")
    for table in sorted(table_summaries.keys()):
        data = table_summaries[table]
        summary_logger.info("Table '%s': col_count_DB_A=%d, col_count_DB_B=%d, row_count_A=%d, row_count_B=%d, rows_only_in_A=%d, rows_only_in_B=%d, rows_mismatched=%d, col_mismatch=%d, checksum_A=%s, checksum_B=%s",
            table,
            data.get('col_count_A', 0),
            data.get('col_count_B', 0),
            data.get('row_count_A', 0),
            data.get('row_count_B', 0),
            data.get('rows_only_in_A', 0),
            data.get('rows_only_in_B', 0),
            data.get('rows_mismatched', 0),
            abs(data.get('col_count_A', 0) - data.get('col_count_B', 0)),
            data.get('checksum_A', ""),
            data.get('checksum_B', "")
        )

# ---------------------------
# Function to Process Allowed and Exclude Object Lists from JSON
# ---------------------------
def process_object_lists(params: dict) -> Tuple[Optional[set], Optional[set]]:
    """
    Processes allowed_objects and exclude_tables from the JSON input.
    If an item ends with ".txt", the function attempts to open the file and read its contents.
    
    :param params: Dictionary parsed from JSON input.
    :return: Tuple (allowed_objects, exclude_tables) as sets (or None if not provided).
    """
    # Process allowed_objects
    allowed_objects = params.get("allowed_objects")
    final_allowed_objects = set()
    if allowed_objects is not None:
        for item in allowed_objects:
            if isinstance(item, str) and item.lower().endswith(".txt"):
                try:
                    with open(item, "r", encoding="utf-8") as f:
                        for line in f:
                            line = line.strip()
                            if line:
                                final_allowed_objects.add(line)
                except Exception as e:
                    final_allowed_objects.add(item)
            else:
                final_allowed_objects.add(item)
    if final_allowed_objects:
        allowed_objects = final_allowed_objects
    else:
        allowed_objects = None

    # Process exclude_tables
    exclude_tables = params.get("exclude_tables")
    final_exclude_tables = set()
    if exclude_tables is not None:
        for item in exclude_tables:
            if isinstance(item, str) and item.lower().endswith(".txt"):
                try:
                    with open(item, "r", encoding="utf-8") as f:
                        for line in f:
                            line = line.strip()
                            if line:
                                final_exclude_tables.add(line)
                except Exception as e:
                    final_exclude_tables.add(item)
            else:
                final_exclude_tables.add(item)
    if final_exclude_tables:
        exclude_tables = final_exclude_tables
    else:
        exclude_tables = None

    return allowed_objects, exclude_tables

# ---------------------------
# Flask Web Application Routes and Templates
# ---------------------------
app = Flask(__name__)

# Base CSS and JavaScript for styling and loading overlay
BASE_CSS = """
<style>
  body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; background: linear-gradient(to right, #ece9e6, #ffffff); }
  h1 { color: #333; }
  .container { background-color: #fff; padding: 20px 30px; border-radius: 8px; box-shadow: 0 4px 8px rgba(0,0,0,0.1); margin-bottom: 20px; }
  textarea { width: 100%; padding: 12px; border: 1px solid #ccc; border-radius: 4px; font-size: 14px; }
  input[type="file"] { margin: 10px 0; }
  input[type="submit"] { background-color: #28a745; color: #fff; padding: 12px 24px; border: none; border-radius: 4px; font-size: 16px; cursor: pointer; }
  input[type="submit"]:hover { background-color: #218838; }
  #loadingOverlay {
    position: fixed;
    top: 0; left: 0;
    width: 100%; height: 100%;
    background-color: rgba(0,0,0,0.5);
    display: none;
    flex-direction: column;
    justify-content: center;
    align-items: center;
    color: #fff;
    font-size: 28px;
    z-index: 1000;
  }
  #loadingOverlay span { font-size: 48px; margin-top: 20px; }
  pre { white-space: pre-wrap; word-wrap: break-word; background-color: #f8f9fa; padding: 15px; border: 1px solid #ddd; border-radius: 4px; }
  a.button { display: inline-block; padding: 10px 20px; background-color: #007bff; color: #fff; text-decoration: none; border-radius: 4px; margin: 5px; }
  a.button:hover { background-color: #0056b3; }
</style>
<script>
  var progressInterval;
  function startProgress() {
    var percent = 0;
    var percentElem = document.getElementById("loadingPercent");
    progressInterval = setInterval(function() {
      if (percent < 99) { percent += 3; } else { percent = 99; }
      percentElem.textContent = percent + "%";
    }, 100);
  }
  function handleSubmit(event) {
    document.querySelector('.container').style.pointerEvents = 'none';
    document.getElementById("loadingOverlay").style.display = "flex";
    startProgress();
  }
  window.onload = function() {
    clearInterval(progressInterval);
    document.getElementById("loadingOverlay").style.display = "none";
  }
</script>
"""

INDEX_TEMPLATE = BASE_CSS + """
<div class="container">
  <h1>Database Comparison Tool</h1>
  <p>Please paste your JSON input below or upload a JSON file.</p>
  <form id="compareForm" method="post" action="/compare" enctype="multipart/form-data" onsubmit="handleSubmit(event)">
    <label for="json_input">Paste JSON input:</label><br>
    <textarea id="json_input" name="json_input" rows="15" placeholder='{
  "dbA": { "host": "", "user": "", "password": "yourpassword", "db": "", "port":  },
  "dbB": { "host": "", "user": "", "password": "yourpassword", "db": "", "port":  },
  "allowed_objects": ["allowed.txt"],
  "exclude_tables": ["table3", "table4"],
  "verbose": false,
  "output_folder": null
}'></textarea><br><br>
    <label for="json_file">Or upload JSON file:</label>
    <input type="file" id="json_file" name="json_file"><br><br>
    <input type="submit" value="Compare Databases">
  </form>
</div>
<div id="loadingOverlay">
  <div>Loading, please wait...</div>
  <span id="loadingPercent">0%</span>
</div>
"""

ERROR_TEMPLATE = BASE_CSS + """
<div class="container">
  <h1>Error</h1>
  <p style="color:#d9534f; font-weight: bold;">{{ error_message }}</p>
  <p>Please check your JSON input. Ensure there are no trailing commas and all property names are enclosed in double quotes.</p>
  <a href="/" class="button">Return to Form</a>
</div>
"""

REPORT_TEMPLATE = BASE_CSS + """
<div class="container">
  <h1>Database Comparison Report</h1>
  <pre>{{ report_content }}</pre>
  <h2>Download Reports</h2>
  <p>You can download the following generated files:</p>
  <a href="/download?folder={{ folder }}&file={{ schema_log_basename }}" class="button">Download Schema Log</a>
  <a href="/download?folder={{ folder }}&file={{ all_log_basename }}" class="button">Download All Log</a>
  <a href="/download?folder={{ folder }}&file={{ summary_log_basename }}" class="button">Download Summary Log</a>
  <a href="/download?folder={{ folder }}&file={{ excel_summary_basename }}" class="button">Download Excel Summary</a>
  <a href="/download?folder={{ folder }}&file={{ word_summary_basename }}" class="button">Download Word Summary</a>
  <br><br>
  <a href="/" class="button">Return to Form</a>
</div>
"""

# ---------------------------
# Flask Route: Index
# ---------------------------
@app.route("/", methods=["GET"])
def index():
    return render_template_string(INDEX_TEMPLATE)

# ---------------------------
# Flask Route: Compare Databases
# ---------------------------
@app.route("/compare", methods=["POST"])
def compare_databases():
    # Retrieve JSON input from file or textarea.
    if "json_file" in request.files and request.files["json_file"].filename:
        file = request.files["json_file"]
        try:
            json_data = file.read().decode("utf-8")
        except Exception as e:
            return render_template_string(ERROR_TEMPLATE, error_message=f"Error reading uploaded file: {e}"), 400
    else:
        json_data = request.form.get("json_input")
        if not json_data:
            return render_template_string(ERROR_TEMPLATE, error_message="No JSON input provided"), 400

    try:
        import json
        params = json.loads(json_data)
    except Exception as e:
        return render_template_string(ERROR_TEMPLATE, error_message=f"Error parsing JSON: {e}"), 400

    # Process allowed_objects and exclude_tables from JSON input.
    # If an item ends with ".txt", read its contents from file.
    allowed_objects, exclude_tables = process_object_lists(params)

    # Setup output folder and filenames.
    output_folder = params.get("output_folder", None)
    output_folder = setup_output_folder(output_folder)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    schema_log_filename = os.path.join(output_folder, f"schema_log_{timestamp}.log")
    all_log_filename = os.path.join(output_folder, f"all_log_{timestamp}.log")
    summary_log_filename = os.path.join(output_folder, f"summary_log_{timestamp}.log")
    excel_summary_filename = os.path.join(output_folder, f"table_summary_{timestamp}.xlsx")
    word_summary_filename = os.path.join(output_folder, f"summary_report_{timestamp}.docx")

    # Initialize summary and table_summaries dictionaries.
    summary: Dict[str, Any] = {
        'tables_only_in_A': 0, 'tables_only_in_B': 0, 'tables_in_both': 0,
        'columns_compared': 0, 'columns_only_in_A': 0, 'columns_only_in_B': 0,
        'indexes_only_in_A': 0, 'indexes_only_in_B': 0, 'indexes_in_both': 0, 'indexes_differ': 0,
        'triggers_only_in_A': 0, 'triggers_only_in_B': 0, 'triggers_in_both': 0, 'triggers_differ': 0,
        'routines_only_in_A': 0, 'routines_only_in_B': 0, 'routines_in_both': 0, 'routines_differ': 0,
        'constraints_only_in_A': 0, 'constraints_only_in_B': 0, 'constraints_in_both': 0, 'constraints_differ': 0,
        'views_only_in_A': 0, 'views_only_in_B': 0, 'views_in_both': 0, 'views_differ': 0,
        'total_tables_compared': 0, 'total_rows_compared': 0,
        'tables_skipped_no_pk': 0, 'rows_only_in_A': 0, 'rows_only_in_B': 0, 'rows_mismatched': 0,
        'schema_diff_tables': [],
        'metadata_diff_tables': [],
        'row_diff_details': {}
    }
    table_summaries: Dict[str, Dict[str, Any]] = {}

    # Get database connection parameters.
    dbA_params = params.get("dbA")
    dbB_params = params.get("dbB")
    verbose = params.get("verbose", False)

    dbA_name = dbA_params.get("db", "DB_A")
    dbB_name = dbB_params.get("db", "DB_B")
    summary['dbA_name'] = dbA_name
    summary['dbB_name'] = dbB_name

    # Setup loggers.
    schema_logger = setup_logger("schemaLogger", schema_log_filename, verbose=verbose)
    all_logger = setup_logger("allLogger", all_log_filename, verbose=verbose)
    summary_logger = setup_logger("summaryLogger", summary_log_filename, verbose=verbose)

    try:
        # Connect to both databases.
        connA = get_mysql_connection(host=dbA_params.get("host"), user=dbA_params.get("user"),
                                     password=dbA_params.get("password"), db=dbA_name,
                                     port=dbA_params.get("port", 3306))
        connB = get_mysql_connection(host=dbB_params.get("host"), user=dbB_params.get("user"),
                                     password=dbB_params.get("password"), db=dbB_name,
                                     port=dbB_params.get("port", 3306))

        # Retrieve and filter schema.
        schemaA = get_tables_and_columns(connA, dbA_name)
        schemaB = get_tables_and_columns(connB, dbB_name)
        if allowed_objects is not None:
            schemaA = {t: cols for t, cols in schemaA.items() if t in allowed_objects}
            schemaB = {t: cols for t, cols in schemaB.items() if t in allowed_objects}
        if exclude_tables is not None:
            schemaA = {t: cols for t, cols in schemaA.items() if t not in exclude_tables}
            schemaB = {t: cols for t, cols in schemaB.items() if t not in exclude_tables}
        compare_schemas(schema_logger, schemaA, schemaB, dbA_name, dbB_name, summary, table_summaries)

        # Retrieve and filter metadata.
        metadataA = get_table_metadata(connA, dbA_name)
        metadataB = get_table_metadata(connB, dbB_name)
        if exclude_tables is not None:
            metadataA = {t: m for t, m in metadataA.items() if t not in exclude_tables}
            metadataB = {t: m for t, m in metadataB.items() if t not in exclude_tables}
        compare_table_metadata(schema_logger, metadataA, metadataB, dbA_name, dbB_name, summary)

        # Retrieve and filter indexes.
        indexesA = get_indexes(connA, dbA_name)
        indexesB = get_indexes(connB, dbB_name)
        if allowed_objects is not None:
            indexesA = {k: v for k, v in indexesA.items() if k[0] in allowed_objects}
            indexesB = {k: v for k, v in indexesB.items() if k[0] in allowed_objects}
        if exclude_tables is not None:
            indexesA = {k: v for k, v in indexesA.items() if k[0] not in exclude_tables}
            indexesB = {k: v for k, v in indexesB.items() if k[0] not in exclude_tables}
        compare_indexes(schema_logger, indexesA, indexesB, dbA_name, dbB_name, summary)

        # Retrieve and filter triggers.
        triggersA = get_triggers(connA, dbA_name)
        triggersB = get_triggers(connB, dbB_name)
        if allowed_objects is not None:
            triggersA = {k: v for k, v in triggersA.items() if v.get('table') in allowed_objects}
            triggersB = {k: v for k, v in triggersB.items() if v.get('table') in allowed_objects}
        if exclude_tables is not None:
            triggersA = {k: v for k, v in triggersA.items() if v.get('table') not in exclude_tables}
            triggersB = {k: v for k, v in triggersB.items() if v.get('table') not in exclude_tables}
        compare_triggers(schema_logger, triggersA, triggersB, dbA_name, dbB_name, summary)

        # Retrieve and filter routines.
        routinesA = get_routines(connA, dbA_name)
        routinesB = get_routines(connB, dbB_name)
        if allowed_objects is not None:
            routinesA = {k: v for k, v in routinesA.items() if k[0] in allowed_objects}
            routinesB = {k: v for k, v in routinesB.items() if k[0] in allowed_objects}
        compare_routines(schema_logger, routinesA, routinesB, dbA_name, dbB_name, summary)

        # Retrieve and filter constraints.
        constraintsA = get_constraints(connA, dbA_name)
        constraintsB = get_constraints(connB, dbB_name)
        if allowed_objects is not None:
            constraintsA = {k: v for k, v in constraintsA.items() if k[0] in allowed_objects}
            constraintsB = {k: v for k, v in constraintsB.items() if k[0] in allowed_objects}
        if exclude_tables is not None:
            constraintsA = {k: v for k, v in constraintsA.items() if k[0] not in exclude_tables}
            constraintsB = {k: v for k, v in constraintsB.items() if k[0] not in exclude_tables}
        compare_constraints(schema_logger, constraintsA, constraintsB, dbA_name, dbB_name, summary)

        # Retrieve and filter views.
        viewsA = get_views(connA, dbA_name)
        viewsB = get_views(connB, dbB_name)
        if allowed_objects is not None:
            viewsA = {k: v for k, v in viewsA.items() if k in allowed_objects}
            viewsB = {k: v for k, v in viewsB.items() if k in allowed_objects}
        if exclude_tables is not None:
            viewsA = {k: v for k, v in viewsA.items() if k not in exclude_tables}
            viewsB = {k: v for k, v in viewsB.items() if k not in exclude_tables}
        compare_views(schema_logger, viewsA, viewsB, dbA_name, dbB_name, summary)

        schema_logger.info(">> Schema Comparison Completed.")

        # Compare row-level data for all common tables.
        compare_all_tables_data(all_logger, connA, connB, dbA_name, dbB_name, schemaA, schemaB, summary, table_summaries)
        all_logger.info(">> Data Comparison Completed.")

        generate_summary_log(summary_logger, summary, table_summaries)
        export_table_summary_to_excel(table_summaries, excel_summary_filename, dbA_name, dbB_name)
        export_summary_to_word(summary, table_summaries, word_summary_filename, dbA_name, dbB_name)
        brief = generate_brief_summary(summary)
        
        schema_log_basename = os.path.basename(schema_log_filename)
        all_log_basename = os.path.basename(all_log_filename)
        summary_log_basename = os.path.basename(summary_log_filename)
        excel_summary_basename = os.path.basename(excel_summary_filename)
        word_summary_basename = os.path.basename(word_summary_filename)
        
        return render_template_string(REPORT_TEMPLATE,
            report_content=brief,
            schema_log=schema_log_filename,
            all_log=all_log_filename,
            summary_log=summary_log_filename,
            excel_summary=excel_summary_filename,
            word_summary=word_summary_filename,
            folder=output_folder,
            schema_log_basename=schema_log_basename,
            all_log_basename=all_log_basename,
            summary_log_basename=summary_log_basename,
            excel_summary_basename=excel_summary_basename,
            word_summary_basename=word_summary_basename)
    except Exception as e:
        return render_template_string(ERROR_TEMPLATE, error_message=f"An error occurred during comparison: {e}"), 500
    finally:
        connA.close()
        connB.close()

# ---------------------------
# Flask Route: Download File
# ---------------------------
@app.route("/download")
def download_file():
    folder = request.args.get("folder")
    file = request.args.get("file")
    if folder and file:
        return send_from_directory(folder, file, as_attachment=True)
    return "Invalid download request", 400

# ---------------------------
# Run Flask Application
# ---------------------------
if __name__ == '__main__':
    app.run(host="0.0.0.0", port=5000)
